"""Build the complete Cell 1 DOCX report from audited source artifacts."""

from __future__ import annotations

import csv
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "reports/23-04-18C-04-cell-1_COMPLETE_MODEL_REPORT.md"
OUTPUT = ROOT / "reports/23-04-18C-04-cell-1_COMPLETE_MODEL_REPORT.docx"

# Selected preset: compact_reference_guide.
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F2F4F7"
MUTED = "666666"
RISK = "9B1C1C"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


FIGURES = {
    "4. Morphology QA": [
        ("figures/morphology/morphology_xy.png", "Figure 1. Native deposited morphology in XY; axon shown separately from soma/dendrites."),
        ("figures/morphology/morphology_xz.png", "Figure 2. Native morphology in XZ, showing the limited sagittal-section depth."),
        ("figures/morphology/morphology_yz.png", "Figure 3. Native morphology in YZ."),
        ("figures/morphology/soma_proximal_zoom_xy.png", "Figure 4. Soma and proximal-neurite zoom; no synthetic process is present."),
    ],
    "8. Passive validation": [
        ("figures/passive/hyperpolarizing_response.png", "Figure 5. Passive response to the standardized -20 pA, 800-ms current step."),
        ("figures/passive/rin_tau_fit.png", "Figure 6. Somatic charging response, mono-exponential tau fit, and Rin probe."),
        ("figures/passive/measured_vs_target.png", "Figure 7. Primary Rin target and non-equivalent capacitance comparisons."),
        ("figures/passive/passive_fitting_error.png", "Figure 8. Retained restrained passive-parameter grid."),
        ("figures/passive/passive_compartment_schematic.png", "Figure 9. Full native-morphology passive compartment scheme."),
    ],
    "12. Active validation": [
        ("figures/active/representative_voltage_traces.png", "Figure 10. Representative current-clamp traces, including strong-current block."),
        ("figures/active/rheobase_trace.png", "Figure 11. Model rheobase response on the final 10-pA validation grid."),
        ("figures/active/fi_curve.png", "Figure 12. F-I curve over the standardized 800-ms protocol."),
        ("figures/active/first_action_potential.png", "Figure 13. First action potential at model rheobase."),
        ("figures/active/ap_phase_plot.png", "Figure 14. First-spike phase plot."),
        ("figures/active/adaptation_isi.png", "Figure 15. Interspike intervals at 120 pA."),
        ("figures/active/experimental_target_vs_model.png", "Figure 16. Signed target ratios exposing active-model mismatches."),
    ],
    "13. Temperature": [
        ("figures/final/temperature_comparison.png", "Figure 17. Primary 23 C model and explicitly predictive 35 C translation at 120 pA."),
    ],
    "15. Numerical robustness": [
        ("figures/final/robustness_summary.png", "Figure 18. Spike-count sensitivity to numerical and model-definition changes."),
    ],
    "18. What is directly supported": [
        ("figures/final/evidence_confidence_schematic.png", "Figure 19. Evidence confidence decreases from same-cell identity to model prediction."),
    ],
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths: list[int], total: int = CONTENT_DXA) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def font_run(run, size=11, bold=None, color=None, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    font_run(run, size=9, color=MUTED)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False


def configure_running_header(header) -> None:
    p = header.paragraphs[0]
    p.text = "NMO_170087 | Complete single-cell model report"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        font_run(run, size=9, color=MUTED)


def configure_running_footer(footer) -> None:
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = p.add_run("NeuropathicPain_Model | Page ")
    font_run(label, size=9, color=MUTED)
    add_page_number(p)


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False
    configure_running_header(section.header)
    configure_running_header(section.even_page_header)
    configure_running_footer(section.footer)
    configure_running_footer(section.even_page_footer)


def cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(82)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("COMPUTATIONAL NEUROSCIENCE MODEL REPORT")
    font_run(run, size=10.5, bold=True, color="7A5A00")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("23-04-18C-04-cell-1")
    font_run(run, size=28, bold=True, color="203748")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run("NMO_170087 mouse PV+/Pax2+ inhibitory interneuron")
    font_run(run, size=14, color="2B5163")
    rows = [
        ("Region", "Lumbar spinal dorsal horn, lamina II-III"),
        ("Medlock mapping", "iPV"),
        ("NEURON", "9.0.1"),
        ("Report date", "2026-08-13"),
        ("Final status", "ENGINEERING READY / BIOLOGICALLY PROVISIONAL"),
        ("Network integration", "NO - unresolved active and robustness failures"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_FILL)
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
        if label == "Network integration":
            for run in cells[1].paragraphs[0].runs:
                run.font.color.rgb = RGBColor.from_string(RISK)
                run.bold = True
    set_table_geometry(table, [2700, 6660])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    run = p.add_run("A traceable, uncertainty-explicit single-cell package for supervisor review and future refinement")
    font_run(run, size=10, italic=True, color=MUTED)


def clean_inline(text: str) -> str:
    text = text.replace("*", "").replace("`", "")
    text = re.sub(r"\[([^]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    return text.strip()


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = [[clean_inline(cell.strip()) for cell in line.strip().strip("|").split("|")] for line in lines]
    if len(rows) >= 2 and all(set(cell) <= {"-", ":"} for cell in rows[1]):
        rows.pop(1)
    columns = len(rows[0])
    table = doc.add_table(rows=0, cols=columns)
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    font_run(run, size=8.2 if columns >= 5 else 9)
            if row_index == 0:
                set_cell_shading(cell, HEADER_FILL)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
        if row_index == 0:
            set_repeat_header(table.rows[-1])
            for cell in table.rows[-1].cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True
    equal = CONTENT_DXA // columns
    widths = [equal] * columns
    widths[-1] += CONTENT_DXA - sum(widths)
    if columns == 2:
        widths = [2800, 6560]
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc: Document, relative: str, caption: str) -> None:
    path = ROOT / relative
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    with Image.open(path) as source:
        ratio = source.height / source.width
    if ratio > 1.05:
        inline = run.add_picture(str(path), height=Inches(6.35))
    else:
        inline = run.add_picture(str(path), width=Inches(6.15))
    inline._inline.docPr.set("descr", caption)
    cp = doc.add_paragraph(caption, style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def parse_markdown(doc: Document) -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    start = next(index for index, line in enumerate(lines) if line.startswith("## 1. Executive summary"))
    index = start
    while index < len(lines):
        line = lines[index].rstrip()
        if not line:
            index += 1
            continue
        if line.startswith("## "):
            title = clean_inline(line[3:])
            heading = doc.add_heading(title, level=1)
            if title.startswith(("1.", "3.", "6.", "9.", "13.", "16.", "17.", "20.", "21.", "22.", "23.")):
                heading.paragraph_format.page_break_before = True
            for relative, caption in FIGURES.get(title, []):
                add_figure(doc, relative, caption)
            index += 1
            continue
        if line.startswith("### "):
            doc.add_heading(clean_inline(line[4:]), level=2)
            index += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(doc, table_lines)
            continue
        if line.startswith("```"):
            index += 1
            code_lines = []
            while index < len(lines) and not lines[index].startswith("```"):
                code_lines.append(lines[index])
                index += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(8)
            set_cell = OxmlElement("w:shd")
            set_cell.set(qn("w:fill"), "F4F6F9")
            p._p.get_or_add_pPr().append(set_cell)
            run = p.add_run("\n".join(code_lines))
            font_run(run, size=8, name="Consolas")
            index += 1
            continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            p.add_run(clean_inline(re.sub(r"^\d+\. ", "", line)))
            index += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(clean_inline(line[2:]))
            index += 1
            continue
        paragraph_lines = [line]
        index += 1
        while index < len(lines) and lines[index].strip() and not lines[index].startswith(("##", "###", "|", "```", "- ")) and not re.match(r"^\d+\. ", lines[index]):
            paragraph_lines.append(lines[index].strip())
            index += 1
        p = doc.add_paragraph(clean_inline(" ".join(paragraph_lines)))
        p.paragraph_format.widow_control = True


def add_code_appendix(doc: Document, title: str, path: Path) -> None:
    doc.add_heading(title, level=1)
    text = path.read_text(encoding="utf-8")
    for chunk_start in range(0, len(text.splitlines()), 48):
        chunk = text.splitlines()[chunk_start : chunk_start + 48]
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F4F6F9")
        p._p.get_or_add_pPr().append(shd)
        run = p.add_run("\n".join(chunk))
        font_run(run, size=6.8, name="Consolas")


def add_csv_table(doc: Document, title: str, path: Path, columns: list[str], labels: list[str] | None = None, font_size=7.2) -> None:
    doc.add_heading(title, level=1)
    with path.open(newline="", encoding="utf-8") as handle:
        rows = list(csv.DictReader(handle))
    labels = labels or columns
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, labels):
        cell.text = label
        set_cell_shading(cell, HEADER_FILL)
        for run in cell.paragraphs[0].runs:
            font_run(run, size=font_size, bold=True)
    set_repeat_header(table.rows[0])
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.keep_with_next = True
    for row in rows:
        cells = table.add_row().cells
        for cell, column in zip(cells, columns):
            cell.text = str(row.get(column, ""))
            for run in cell.paragraphs[0].runs:
                font_run(run, size=font_size)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.0
    equal = CONTENT_DXA // len(columns)
    widths = [equal] * len(columns)
    widths[-1] += CONTENT_DXA - sum(widths)
    set_table_geometry(table, widths)


def add_record_appendix(doc: Document, title: str, path: Path, key_field: str) -> None:
    doc.add_heading(title, level=1)
    with path.open(newline="", encoding="utf-8") as handle:
        rows = list(csv.DictReader(handle))
    for index, row in enumerate(rows, 1):
        heading = clean_inline(row.get(key_field) or f"Record {index}")
        doc.add_heading(f"{title.split('.')[0]}.{index} {heading}", level=2)
        for label, value in row.items():
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1)
            paragraph.paragraph_format.line_spacing = 1.0
            label_run = paragraph.add_run(f"{label}: ")
            font_run(label_run, size=7.5, bold=True, color=DARK_BLUE)
            value_run = paragraph.add_run(value or "")
            font_run(value_run, size=7.5)


def appendices(doc: Document) -> None:
    heading = doc.add_heading("Appendices", level=1)
    heading.paragraph_format.page_break_before = True
    p = doc.add_paragraph("These appendices embed the complete selected configuration and machine-readable validation records. Repository CSV/JSON files remain authoritative for numerical precision.")
    p.paragraph_format.space_after = Pt(10)
    add_code_appendix(doc, "Appendix A. Full final parameter JSON", ROOT / "parameters/final/NMO_170087_final_23C.json")
    add_csv_table(
        doc,
        "Appendix B. Full active validation table",
        ROOT / "results/validation/active_validation_table.csv",
        ["metric", "experimental_target", "model_value", "acceptance_criterion", "status", "evidence_level"],
        ["Metric", "Target", "Model", "Acceptance", "Status", "Evidence"],
        font_size=6.6,
    )
    add_csv_table(
        doc,
        "Appendix C. Full F-I numerical data",
        ROOT / "results/active/FI_curve.csv",
        ["current_na", "spike_count", "frequency_hz", "first_spike_latency_ms", "last_spike_time_ms", "mean_isi_ms", "adaptation_ratio_last_over_first", "tonic_persistence_fraction"],
        ["I (nA)", "Spikes", "Hz", "Latency", "Last spike", "Mean ISI", "Adapt.", "Persistence"],
        font_size=6.4,
    )
    add_code_appendix(doc, "Appendix D. Full morphology QA JSON", ROOT / "results/morphology_qa/morphology_qa.json")
    add_record_appendix(doc, "Appendix E. Full evidence matrix", ROOT / "docs/evidence/evidence_matrix.csv", "Parameter")
    add_record_appendix(doc, "Appendix F. Full channel justification", ROOT / "docs/evidence/channel_justification.csv", "Channel/current")
    add_csv_table(
        doc,
        "Appendix G. Full robustness data",
        ROOT / "results/robustness/robustness_summary.csv",
        ["family", "condition", "nseg_total", "spike_count", "frequency_hz", "tonic_persistence_fraction", "ap_threshold_mv", "ap_peak_mv", "ap_half_width_ms"],
        ["Family", "Condition", "nseg", "Spikes", "Hz", "Persist.", "Threshold", "Peak", "Width"],
        font_size=6.1,
    )


def audit(doc: Document) -> None:
    section = doc.sections[0]
    assert abs(section.page_width.inches - 8.5) < 0.01
    assert abs(section.page_height.inches - 11) < 0.01
    assert all(abs(value.inches - 1.0) < 0.01 for value in [section.top_margin, section.right_margin, section.bottom_margin, section.left_margin])
    assert abs(section.header_distance.inches - 0.492) < 0.01
    assert abs(section.footer_distance.inches - 0.492) < 0.01
    assert all(table._tbl.tblPr.first_child_found_in("w:tblW") is not None for table in doc.tables)
    assert len(doc.inline_shapes) >= 19


def main() -> None:
    doc = Document()
    # Explicitly populate both default and even-page stories so Word and
    # LibreOffice render identical running headers and footers.
    doc.settings.odd_and_even_pages_header_footer = True
    configure_styles(doc)
    configure_section(doc.sections[0])
    cover(doc)
    parse_markdown(doc)
    appendices(doc)
    audit(doc)
    doc.core_properties.title = "NMO_170087 complete single-cell model report"
    doc.core_properties.subject = "Mouse PV+/Pax2+ spinal dorsal horn neuron model"
    doc.core_properties.author = "NeuropathicPain_Model project"
    doc.core_properties.keywords = "NEURON, spinal dorsal horn, parvalbumin, NMO_170087, computational neuroscience"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(json.dumps({"output": str(OUTPUT), "paragraphs": len(doc.paragraphs), "tables": len(doc.tables), "figures": len(doc.inline_shapes)}, indent=2))


if __name__ == "__main__":
    main()
