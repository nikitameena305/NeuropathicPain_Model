#!/usr/bin/env python3
"""Build the complete L292-E1-LCN Markdown and Microsoft Word report.

The report is assembled only from files already present in the isolated
exc_interneuron workspace.  It intentionally preserves the failed delayed
35 °C validation gate and excludes the unrelated 14-1-15-A-A2sep neuron.
"""

from __future__ import annotations

import csv
import hashlib
import json
import math
import os
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
FIG_DIR = REPORTS / "L292_E1_LCN_report_figures"
DOCX_PATH = REPORTS / "L292_E1_LCN_COMPLETE_EXCITATORY_INTERNEURON_REPORT.docx"
MD_PATH = REPORTS / "L292_E1_LCN_COMPLETE_EXCITATORY_INTERNEURON_REPORT_source.md"

NAVY = "12324A"
BLUE = "2878A5"
TEAL = "1B8A89"
GREEN = "2E7D5B"
GOLD = "B78324"
RED = "A83232"
GRAY = "5E6A71"
LIGHT_GRAY = "E9EEF1"
PALE_BLUE = "E8F1F6"
PALE_GREEN = "E7F4EC"
PALE_RED = "F9E8E8"
PALE_GOLD = "F8F0DC"
WHITE = "FFFFFF"
BLACK = "1A1A1A"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def load_json(rel: str):
    with (ROOT / rel).open("r", encoding="utf-8") as fh:
        return json.load(fh)


def read_csv(rel: str):
    with (ROOT / rel).open("r", newline="", encoding="utf-8") as fh:
        rows = []
        for raw in csv.DictReader(fh):
            row = {}
            for key, value in raw.items():
                if value == "":
                    row[key] = None
                elif value in {"True", "False"}:
                    row[key] = value == "True"
                else:
                    try:
                        row[key] = float(value)
                    except (TypeError, ValueError):
                        row[key] = value
            rows.append(row)
        return rows


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def fmt(value, digits=3, missing="—"):
    if value is None:
        return missing
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, int):
        return f"{value:,}"
    if isinstance(value, float):
        if math.isnan(value):
            return missing
        return f"{value:.{digits}f}"
    return str(value)


def set_run_font(run, name="Aptos", size=None, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="B9C5CB", size=5):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_toc_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and update field if page numbers are not refreshed."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])


def add_hyperlink(paragraph, text: str, target: str):
    part = paragraph.part
    rel_id = part.relate_to(target, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                            is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([r_pr, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Title": (30, NAVY, 0, 10),
        "Subtitle": (14, BLUE, 0, 12),
        "Heading 1": (17, NAVY, 16, 8),
        "Heading 2": (13.5, BLUE, 12, 6),
        "Heading 3": (11.5, TEAL, 9, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Normal" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:ascii"), style.font.name)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = styles["Caption"]
    caption.font.name = "Aptos"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    caption.font.size = Pt(8.5)
    caption.font.color.rgb = rgb(GRAY)
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(9)
    caption.paragraph_format.keep_together = True

    for name, indent, hanging in (("List Bullet", 0.375, 0.194), ("List Number", 0.375, 0.194)):
        style = styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(indent)
        style.paragraph_format.first_line_indent = Inches(-hanging)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    code = styles.add_style("Report Code", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(7.7)
    code.font.color.rgb = rgb(NAVY)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.left_indent = Inches(0.14)


class ReportWriter:
    def __init__(self, doc: Document):
        self.doc = doc
        self.md: list[str] = []
        self.figure_count = 0
        self.table_count = 0

    def h1(self, text):
        self.doc.add_heading(text, level=1)
        self.md += [f"# {text}", ""]

    def h2(self, text):
        self.doc.add_heading(text, level=2)
        self.md += [f"## {text}", ""]

    def h3(self, text):
        self.doc.add_heading(text, level=3)
        self.md += [f"### {text}", ""]

    def para(self, text, bold_prefix=None, italic=False):
        p = self.doc.add_paragraph()
        if bold_prefix and text.startswith(bold_prefix):
            r = p.add_run(bold_prefix)
            set_run_font(r, bold=True)
            r2 = p.add_run(text[len(bold_prefix):])
            set_run_font(r2, italic=italic)
        else:
            r = p.add_run(text)
            set_run_font(r, italic=italic)
        self.md += [text, ""]
        return p

    def bullets(self, items: Iterable[str]):
        for item in items:
            p = self.doc.add_paragraph(style="List Bullet")
            set_run_font(p.add_run(item))
            self.md.append(f"- {item}")
        self.md.append("")

    def numbered(self, items: Iterable[str]):
        for idx, item in enumerate(items, 1):
            p = self.doc.add_paragraph(style="List Number")
            set_run_font(p.add_run(item))
            self.md.append(f"{idx}. {item}")
        self.md.append("")

    def callout(self, label: str, text: str, kind="info"):
        color, fill = {
            "pass": (GREEN, PALE_GREEN), "fail": (RED, PALE_RED),
            "warning": (GOLD, PALE_GOLD), "info": (BLUE, PALE_BLUE),
        }[kind]
        table = self.doc.add_table(rows=1, cols=1)
        set_table_geometry(table, [9360], indent_dxa=120)
        set_table_borders(table, color=color, size=10)
        cell = table.cell(0, 0)
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(f"{label}: ")
        set_run_font(r, size=10.5, color=color, bold=True)
        r = p.add_run(text)
        set_run_font(r, size=10.2, color=NAVY, bold=(kind == "fail"))
        self.doc.add_paragraph().paragraph_format.space_after = Pt(1)
        self.md += [f"> **{label}:** {text}", ""]

    def table(self, caption: str, headers: Sequence[str], rows: Sequence[Sequence], widths: Sequence[int],
              font_size=8.2, status_col=None, manifest=False):
        self.table_count += 1
        cap = self.doc.add_paragraph(style="Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(cap.add_run(f"Table {self.table_count}. {caption}"), size=8.5, color=GRAY,
                     italic=True)
        table = self.doc.add_table(rows=1, cols=len(headers))
        set_table_geometry(table, widths, indent_dxa=120)
        set_table_borders(table)
        table.rows[0]._tr.get_or_add_trPr()
        set_repeat_table_header(table.rows[0])
        for idx, (cell, header) in enumerate(zip(table.rows[0].cells, headers)):
            set_cell_shading(cell, NAVY)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(str(header)), size=font_size, color=WHITE, bold=True)
        for ridx, row in enumerate(rows):
            cells = table.add_row().cells
            if ridx % 2:
                for cell in cells:
                    set_cell_shading(cell, "F7F9FA")
            for idx, (cell, value) in enumerate(zip(cells, row)):
                set_cell_margins(cell, top=75 if manifest else 90, bottom=75 if manifest else 90)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05 if manifest else 1.12
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (idx > 0 and len(str(value)) < 24) else WD_ALIGN_PARAGRAPH.LEFT
                color = BLACK
                bold = False
                if status_col is not None and idx == status_col:
                    upper = str(value).upper()
                    if "FAIL" in upper or "NOT READY" in upper:
                        color, bold = RED, True
                    elif "PASS" in upper or "READY" in upper or "VALIDATED" in upper:
                        color, bold = GREEN, True
                    elif "GATED" in upper or "UNKNOWN" in upper:
                        color, bold = GOLD, True
                set_run_font(p.add_run(str(value)), name="Consolas" if manifest and idx in {0, 2} else "Aptos",
                             size=font_size, color=color, bold=bold)
        self.doc.add_paragraph().paragraph_format.space_after = Pt(1)
        self.md += [f"**Table {self.table_count}. {caption}**", ""]
        self.md.append("| " + " | ".join(headers) + " |")
        self.md.append("|" + "|".join(["---"] * len(headers)) + "|")
        for row in rows:
            escaped = [str(v).replace("|", "\\|").replace("\n", "<br>") for v in row]
            self.md.append("| " + " | ".join(escaped) + " |")
        self.md.append("")
        return table

    def figure(self, filename: str, caption: str, alt: str, width=6.25):
        path = FIG_DIR / filename
        if not path.exists():
            raise FileNotFoundError(path)
        self.figure_count += 1
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        shape = p.add_run().add_picture(str(path), width=Inches(width))
        shape._inline.docPr.set("descr", alt)
        cap = self.doc.add_paragraph(style="Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(cap.add_run(f"Figure {self.figure_count}. {caption}"), size=8.5, color=GRAY, italic=True)
        self.md += [f"![{alt}](L292_E1_LCN_report_figures/{filename})", "",
                    f"*Figure {self.figure_count}. {caption}*", ""]

    def code(self, text: str):
        p = self.doc.add_paragraph(style="Report Code")
        set_run_font(p.add_run(text), name="Consolas", size=7.7, color=NAVY)
        set_cell = None
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F3F6F8")
        p_pr.append(shd)
        self.md += ["```text", text, "```", ""]

    def link_item(self, number: int, label: str, rel: str, text: str):
        p = self.doc.add_paragraph(style="List Number")
        set_run_font(p.add_run(f"{label}: "), bold=True)
        target = (ROOT / rel).resolve().as_uri()
        add_hyperlink(p, rel, target)
        set_run_font(p.add_run(f" — {text}"))
        self.md.append(f"{number}. **{label}:** [{rel}](../{rel}) — {text}")

    def page_break(self):
        self.doc.add_page_break()
        self.md += ["<div style=\"page-break-after: always;\"></div>", ""]


def setup_document() -> Document:
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    doc.settings.element.append(update)
    return doc


def add_cover(doc: Document):
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    set_run_font(p.add_run("COMPUTATIONAL NEUROSCIENCE MODEL REPORT"), size=10.5, color=GOLD, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    set_run_font(p.add_run("Development and Validation of a Morphology-Constrained Rat Lamina-I\n"
                           "Excitatory Interneuron Model Based on L292-E1-LCN"),
                 name="Aptos Display", size=25, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    set_run_font(p.add_run("NEURON-Based Reconstruction and Intrinsic Modelling of Six Excitatory\n"
                           "Spinal Dorsal Horn Populations for a Neuropathic Pain Circuit"),
                 size=13, color=BLUE, italic=True)

    rows = [
        ("Neuron morphology", "L292-E1-LCN"), ("NeuroMorpho ID", "NMO_34021"),
        ("Species / strain", "Rat / Wistar"), ("Region", "Lumbar spinal dorsal horn, lamina I"),
        ("Morphological class", "Multipolar local-circuit neuron"),
        ("Morphology completeness", "Soma + dendrites + axon"),
        ("Simulation platform", "NEURON 9.0.1"),
        ("Temperature benchmarks", "22–24 °C source; 23 °C validation; 35 °C target"),
        ("Report date", "12 August 2026"),
    ]
    table = doc.add_table(rows=0, cols=2)
    set_table_geometry(table, [2700, 6660], indent_dxa=120)
    set_table_borders(table, color="CCD6DB", size=4)
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], PALE_BLUE)
        for c in cells:
            set_cell_margins(c, top=85, bottom=85)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_run_font(cells[0].paragraphs[0].add_run(label), size=8.8, color=NAVY, bold=True)
        set_run_font(cells[1].paragraphs[0].add_run(value), size=8.8, color=BLACK)

    doc.add_paragraph()
    status = doc.add_table(rows=1, cols=1)
    set_table_geometry(status, [9360], indent_dxa=120)
    set_table_borders(status, color=RED, size=14)
    cell = status.cell(0, 0)
    set_cell_shading(cell, PALE_RED)
    set_cell_margins(cell, top=180, bottom=180, start=220, end=220)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("CURRENT MODEL STATUS: NOT READY\n"), size=18, color=RED, bold=True)
    set_run_font(p.add_run("Main failed gate: the common delayed-excitatory intrinsic model develops "
                           "depolarization block at 35 °C under moderate-to-strong depolarizing current."),
                 size=10.5, color=NAVY, bold=True)


def configure_body_section(doc: Document):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.82)
    sec.bottom_margin = Inches(0.78)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.38)
    sec.footer_distance = Inches(0.42)
    sec.header.is_linked_to_previous = False
    sec.footer.is_linked_to_previous = False
    header = sec.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run("L292-E1-LCN  |  Complete excitatory-interneuron report"),
                 size=8.2, color=GRAY, bold=True)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "CBD5DA")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    footer = sec.footer
    add_page_field(footer.paragraphs[0])
    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:start"), "1")
    sec._sectPr.append(pg_num_type)


def file_purpose(rel: str) -> tuple[str, str]:
    low = rel.lower().replace("\\", "/")
    if low.endswith(".swc"):
        return "Morphology input", "DIRECT INPUT"
    if "/morphology/" in f"/{low}" or low.startswith("morphology/"):
        return "Morphology provenance", "PROVENANCE"
    if low.endswith(".mod"):
        return "NMODL mechanism source", "STAGED SOURCE"
    if "mechanisms/x86_64" in low or "__pycache__" in low:
        return "Derived build artifact", "EXCLUDED"
    if low.endswith(".json") and low.startswith("parameters/"):
        return "Model configuration/reference", "CONFIGURATION"
    if low.startswith("results/"):
        status = "DIAGNOSTIC"
        if any(x in low for x in ("/final/", "final_strict_dlambda", "final_after_hh2")):
            status = "ACCEPTED RESULT"
        if "convergence/" in low and "convergence_strict" not in low:
            status = "SUPERSEDED RESULT"
        return "Simulation/analysis output", status
    if low.startswith("scripts/"):
        return "Reproducibility script", "SOURCE"
    if low.startswith("docs/"):
        return "Scientific audit/documentation", "DOCUMENTATION"
    if low.startswith("reports/l292_e1_lcn_report_figures/"):
        return "Report figure", "GENERATED FIGURE"
    if low.startswith("reports/"):
        return "Project report/log", "REPORT"
    return "Project file", "PROJECT"


def build_manifest():
    excluded_parts = {"__pycache__", "x86_64"}
    excluded_names = {DOCX_PATH.name, MD_PATH.name}
    rows = []
    for path in sorted(ROOT.rglob("*")):
        if not path.is_file() or any(part in excluded_parts for part in path.parts):
            continue
        if path.name in excluded_names:
            continue
        rel = path.relative_to(ROOT).as_posix()
        purpose, status = file_purpose(rel)
        rows.append((rel, purpose, sha256(path), status))
    return rows
